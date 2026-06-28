"""Generate a Word .docx project summary for Local RAG Chat."""
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

# Title
title = doc.add_heading("Project Summary — Local RAG Chat", level=0)
sub = doc.add_paragraph()
run = sub.add_run("Repository: https://github.com/saidataharimatohoku-max/local-rag-chat")
run.italic = True

doc.add_heading("Overview", level=1)
doc.add_paragraph(
    "Local RAG Chat is a Retrieval-Augmented Generation (RAG) application that "
    "answers questions grounded in your own documents and cites its sources, so "
    "it does not invent answers. It runs fully local and free with Ollama — no "
    "cloud account required — and can optionally run on Azure OpenAI + Azure AI "
    "Search by changing configuration only (no code changes)."
)

doc.add_heading("What it does", level=1)
for line in [
    "You add documents (Markdown, plain text, PDF, or Word) — from the browser or the data/ folder.",
    "The app splits each document into overlapping chunks and turns them into vector embeddings stored in a searchable index.",
    "When you ask a question, it embeds the question, finds the most relevant chunks, and asks a language model to answer using only that context.",
    "The answer streams back token-by-token, with each cited source expandable to reveal the exact text it was based on.",
]:
    doc.add_paragraph(line, style="List Number")

doc.add_heading("Key features", level=1)
for feat in [
    "Grounded answers with citations — each cited source is an expandable snippet of the retrieved text.",
    "Streaming responses — answers render token-by-token via Server-Sent Events.",
    "Multi-format ingestion — Markdown (.md), text (.txt), PDF (.pdf), and Word (.docx).",
    "In-browser upload — add a document with the + Add document button or by dragging and dropping it onto the page; it is indexed immediately and the indexed-document list refreshes.",
    "Local-first and free — Ollama with llama3.2:1b (chat) and nomic-embed-text (embeddings); a NumPy cosine-similarity store replaces a cloud vector database.",
    "Cloud-ready — the same code targets Azure OpenAI and Azure AI Search via environment variables; Bicep templates included.",
    "Tested + CI — a pytest suite runs offline (the model is mocked) and on every push via GitHub Actions.",
]:
    doc.add_paragraph(feat, style="List Bullet")

doc.add_heading("Architecture", level=1)
arch = (
    "Browser (HTML/CSS/JS)\n"
    "    | POST /api/chat/stream (SSE), POST /api/upload\n"
    "    v\n"
    "FastAPI app (backend/app.py)\n"
    "    v\n"
    "RAG pipeline (backend/rag.py): embed question -> retrieve top-k chunks -> chat model answers\n"
    "    - local: NumPy cosine search (backend/store.py)\n"
    "    - azure: Azure AI Search\n"
    "Ingestion (backend/ingest.py): read -> chunk -> embed -> index"
)
mono = doc.add_paragraph()
run = mono.add_run(arch)
run.font.name = "Consolas"
run.font.size = Pt(9)
doc.add_paragraph(
    "The active provider (local vs. Azure) is selected automatically in "
    "backend/config.py: it uses Azure when Azure OpenAI credentials are present, "
    "otherwise it runs locally with Ollama. The PROVIDER environment variable "
    "can force either mode."
)

doc.add_heading("Tech stack", level=1)
stack = [
    ("Backend", "Python, FastAPI, Uvicorn, Pydantic"),
    ("LLM (local)", "Ollama — llama3.2:1b, nomic-embed-text"),
    ("LLM (cloud)", "Azure OpenAI (chat + embeddings)"),
    ("Retrieval", "Local NumPy vector store, or Azure AI Search"),
    ("Documents", "pypdf, python-docx"),
    ("Frontend", "Vanilla HTML, CSS, JavaScript (Server-Sent Events)"),
    ("Infra", "Azure Bicep (AI Search + Linux Python Web App)"),
    ("Testing / CI", "pytest, FastAPI TestClient, GitHub Actions"),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Layer"
hdr[1].text = "Technology"
for layer, tech in stack:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech

doc.add_heading("Project structure", level=1)
structure = (
    "backend/   FastAPI app, RAG pipeline, ingestion, vector store, config\n"
    "frontend/  Static chat UI (HTML/CSS/JS)\n"
    "data/      Source documents to index (.md, .txt, .pdf, .docx)\n"
    "infra/     Azure Bicep templates (Search + Web App)\n"
    "tests/     Pytest suite (chunking, store, documents, config, API)\n"
    "docs/      Demo screenshot and this summary\n"
    "scripts/   Word-doc generators and helpers\n"
    "backup.ps1 Local backup script (git bundle + source zip)\n"
    ".github/   CI (tests) and deploy workflows"
)
run = doc.add_paragraph().add_run(structure)
run.font.name = "Consolas"
run.font.size = Pt(9)

doc.add_heading("How to run locally (free)", level=1)
steps = (
    "ollama pull llama3.2:1b\n"
    "ollama pull nomic-embed-text\n"
    "python -m venv .venv\n"
    ".venv\\Scripts\\python.exe -m pip install -r requirements.txt\n"
    ".venv\\Scripts\\python.exe -m backend.ingest\n"
    ".venv\\Scripts\\python.exe -m uvicorn backend.app:app --reload\n"
    "# open http://localhost:8000"
)
run = doc.add_paragraph().add_run(steps)
run.font.name = "Consolas"
run.font.size = Pt(9)

doc.add_heading("Testing", level=1)
doc.add_paragraph(
    "The suite (25 tests) covers text chunking, the on-disk vector store, "
    "multi-format document reading, provider selection, and the API endpoints "
    "(chat, streaming, upload, validation, and a path-traversal security check). "
    "It runs fully offline because the language model is mocked."
)

doc.add_heading("Security notes", level=1)
for note in [
    "Secrets (.env) and the generated index (data/index.json) are git-ignored.",
    "The upload endpoint validates file extensions against an allow-list and sanitizes filenames to prevent path-traversal writes outside data/.",
]:
    doc.add_paragraph(note, style="List Bullet")

doc.add_heading("Backup & versioning", level=1)
for note in [
    "Version-controlled with Git and published on GitHub, with tagged releases (e.g. v1.0.0) marking stable, restorable checkpoints.",
    "backup.ps1 creates two local backups in a sibling folder: a git bundle with the complete commit history (clone it offline with git clone <file>.bundle) and a source zip snapshot of tracked files.",
]:
    doc.add_paragraph(note, style="List Bullet")

doc.add_heading("Possible future enhancements", level=1)
for item in [
    "Conversation memory for follow-up questions.",
    "Heading- or paragraph-aware chunking for more precise retrieval.",
    "Dockerfile + docker-compose (app + Ollama) for one-command setup.",
    "Authentication and per-user document collections.",
]:
    doc.add_paragraph(item, style="List Bullet")

out_path = r"c:\Users\T495s\Downloads\microsoft (project 4)\docs\Project_Summary.docx"
doc.save(out_path)
print(f"Saved {out_path}")
