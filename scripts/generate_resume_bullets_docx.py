"""Generate a Word .docx of resume bullets for Local RAG Chat."""
from docx import Document
from docx.shared import Pt

doc = Document()

doc.add_heading("Local RAG Chat — Resume Bullets", level=0)
p = doc.add_paragraph()
r = p.add_run(
    "Project: Local RAG Chat | Python, FastAPI, Ollama, Azure OpenAI | "
    "github.com/saidataharimatohoku-max/local-rag-chat"
)
r.italic = True

doc.add_heading("Full version (Projects entry)", level=1)
full = [
    "Built a full-stack Retrieval-Augmented Generation (RAG) chat application that answers questions grounded in private documents with cited sources, eliminating LLM hallucination by constraining responses to retrieved context.",
    "Engineered a provider-abstracted backend (FastAPI) that runs 100% locally and free via Ollama or scales to Azure OpenAI + Azure AI Search through configuration alone, with zero code changes.",
    "Implemented multi-format ingestion (Markdown, TXT, PDF, DOCX), a NumPy cosine-similarity vector store, and token-by-token streaming responses (Server-Sent Events) for a responsive UX.",
    "Designed an agentic mode that routes each query (answer/search/clarify), uses the LLM to self-evaluate retrieval quality, and rephrases-and-retries the search before answering, surfacing its reasoning steps to the user.",
    "Added in-browser document upload with server-side validation and path-traversal protection, plus a 48-test pytest suite (81% coverage) and GitHub Actions CI that runs fully offline.",
]
for b in full:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Short version (2 lines)", level=1)
short = [
    "Developed a full-stack RAG Q&A app (Python, FastAPI, Ollama/Azure OpenAI) that answers questions from private documents with cited sources, running locally and free or scaling to Azure via config only.",
    "Built an agentic mode (query routing, LLM-based retrieval self-evaluation, and rephrase-and-retry), multi-format ingestion (PDF/DOCX/MD/TXT), streaming responses, secure in-browser upload, and a 48-test suite (81% coverage) with GitHub Actions CI.",
]
for b in short:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("One-line version (summary/skills line)", level=1)
doc.add_paragraph(
    "Built a local-first, agentic RAG document-Q&A app (FastAPI + Ollama, "
    "Azure-ready) with query routing, retrieval self-evaluation, streaming "
    "responses, cited sources, multi-format ingestion, and CI-tested APIs.",
    style="List Bullet",
)

out_path = r"c:\Users\T495s\Downloads\microsoft (project 4)\docs\Resume_Bullets.docx"
doc.save(out_path)
print(f"Saved {out_path}")
