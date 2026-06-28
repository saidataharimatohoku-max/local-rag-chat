"""Tests for multi-format document reading."""
import pytest

from backend import ingest as ingest_module
from backend.ingest import read_document


def test_reads_markdown(tmp_path):
    path = tmp_path / "note.md"
    path.write_text("# Title\n\nSome text.", encoding="utf-8")
    assert "Some text." in read_document(str(path))


def test_reads_plain_text(tmp_path):
    path = tmp_path / "note.txt"
    path.write_text("hello from txt", encoding="utf-8")
    assert read_document(str(path)) == "hello from txt"


def test_reads_docx(tmp_path):
    from docx import Document

    path = tmp_path / "note.docx"
    document = Document()
    document.add_paragraph("first paragraph")
    document.add_paragraph("second paragraph")
    document.save(str(path))

    text = read_document(str(path))
    assert "first paragraph" in text
    assert "second paragraph" in text


def test_unsupported_extension_raises(tmp_path):
    path = tmp_path / "note.csv"
    path.write_text("a,b,c", encoding="utf-8")
    with pytest.raises(ValueError):
        read_document(str(path))


def test_list_documents_filters_supported(monkeypatch, tmp_path):
    (tmp_path / "a.md").write_text("x", encoding="utf-8")
    (tmp_path / "b.csv").write_text("x", encoding="utf-8")
    (tmp_path / "c.docx").write_bytes(b"PK")
    monkeypatch.setattr(ingest_module, "DATA_DIR", str(tmp_path))

    docs = ingest_module.list_documents()
    assert docs == ["a.md", "c.docx"]


def test_ingest_empty_data_returns_zero(monkeypatch, tmp_path):
    monkeypatch.setattr(ingest_module, "DATA_DIR", str(tmp_path))
    assert ingest_module.ingest() == 0
